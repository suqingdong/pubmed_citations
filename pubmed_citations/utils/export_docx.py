import re

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


def add_formatted_paragraph(doc, text, font='Times New Roman', eastFont='宋体', size=12):
    # use regex to find <i>...</i> tags
    parts = re.split(r'(<i>.*?</i>)', text)
    
    # create a new paragraph
    paragraph = doc.add_paragraph()
    
    for part in parts:
        if part.startswith('<i>') and part.endswith('</i>'):
            # remove <i> tags
            content = part[3:-4]
            run = paragraph.add_run(content)
            run.italic = True
        else:
            run = paragraph.add_run(part)
        
        # set Western language font and size
        run.font.name = font
        # set East Asian language font and size
        run._element.rPr.rFonts.set(qn('w:eastAsia'), eastFont)
        run.font.size = Pt(size)


def create_docx(texts, outfile='output.docx'):
    doc = Document()
    
    for text in texts:
        line = '\t'.join(text)
        add_formatted_paragraph(doc, line)
    
    doc.save(outfile)


if __name__ == '__main__':
    pmid = '31171989'
    texts = [
        [
            pmid,
            'Sinha BK, Mason RP. IS METABOLIC ACTIVATION OF TOPOISOMERASE II POISONS IMPORTANT IN THE MECHANISM OF CYTOTOXICITY?. <i>J Drug Metab Toxicol</i>. 2015;6(3):186. doi:10.4172/2157-7609.1000186',
        ],
        [
            pmid,
            'Sinha, B. K., & Mason, R. P. (2015). IS METABOLIC ACTIVATION OF TOPOISOMERASE II POISONS IMPORTANT IN THE MECHANISM OF CYTOTOXICITY?. <i>Journal of drug metabolism & toxicology</i>, <i>6</i>(3), 186. https://doi.org/10.4172/2157-7609.1000186',
        ]
    ]
    create_docx(texts)
